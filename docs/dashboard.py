"""
dashboard.py
============
Competitive Intelligence Dashboard.
Reads from monthly_market_trends_YYYY_MM.json

Run: streamlit run dashboard.py
"""

import glob
import json
import os
from datetime import datetime
from pathlib import Path

import pandas as pd
import plotly.express as px
import streamlit as st

BRIEF_DIR  = os.environ.get("BRIEF_DIR", "data")

# ─────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────

st.set_page_config(
    page_title="Competitive Intelligence",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
    .block-container { padding: 1.5rem 2rem; }
    .exec-box {
        background: linear-gradient(135deg, #1a1a2e 0%, #16213e 100%);
        border-radius: 12px; padding: 1.5rem 2rem;
        border-left: 4px solid #e94560; margin-bottom: 1rem;
        font-size: 0.97rem; line-height: 1.7; color: #ddd;
    }
    .trend-box {
        background: #0f3460; border-radius: 8px;
        padding: 1rem 1.2rem; height: 100%;
        border-left: 3px solid #e94560;
    }
    .trend-title { font-weight: 700; color: #e94560; margin-bottom: 0.4rem; }
    .trend-body  { font-size: 0.88rem; color: #ccc; line-height: 1.5; }
    .comp-header {
        font-size: 1.25rem; font-weight: 700; color: #e94560;
        margin: 2rem 0 0.5rem; border-bottom: 1px solid #333;
        padding-bottom: 0.4rem;
    }
    .comp-summary {
        background: #16213e; border-radius: 8px;
        padding: 1rem 1.2rem; margin-bottom: 0.8rem;
        font-size: 0.93rem; line-height: 1.6; color: #ddd;
    }
    .theme-row {
        background: #0d0d1a; border-radius: 6px;
        padding: 0.7rem 1rem; margin-bottom: 0.35rem;
        border-left: 3px solid #333;
    }
    .theme-summary { color: #eee; font-size: 0.9rem; line-height: 1.5; }
    .theme-meta { color: #777; font-size: 0.78rem; margin-top: 0.25rem; }
    .section-title {
        font-size: 1rem; font-weight: 600; color: #e94560;
        margin: 1.5rem 0 0.8rem; border-left: 3px solid #e94560;
        padding-left: 0.7rem;
    }
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────
# DATA LOADING
# ─────────────────────────────────────────────

@st.cache_data(ttl=300)
def load_brief(brief_dir):
    files = sorted(
        glob.glob(os.path.join(brief_dir, "monthly_market_trends_*.json")),
        reverse=True
    )
    if not files:
        return None, None
    with open(files[0]) as f:
        brief = json.load(f)
    return brief, brief.get("year_month", "")


brief, year_month = load_brief(BRIEF_DIR)

# ─────────────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────────────

with st.sidebar:
    st.markdown("### 🛡️ Competitive Intel")
    st.markdown(f"**{year_month or 'No data yet'}**")
    st.markdown("---")

    selected_companies = []
    if brief:
        all_companies = list(brief.get("competitor_summaries", {}).keys())
        selected_companies = st.multiselect(
            "Filter competitors", all_companies, default=all_companies
        )

    st.markdown("---")
    st.markdown("### 📥 Export")
    export_btn = st.button("📄 Export to Word Doc", use_container_width=True)
    st.markdown("---")
    st.caption("Pipeline: Google News RSS → NLP dedup → DBSCAN → Claude AI → DuckDB")


# ─────────────────────────────────────────────
# NO DATA STATE
# ─────────────────────────────────────────────

if not brief:
    st.markdown("# 🛡️ Competitive Intelligence Dashboard")
    st.warning("No brief data found. Run `generate_brief.py` via GitHub Actions first.")
    st.stop()

# Apply company filter
theme_summaries      = [t for t in brief.get("theme_summaries", [])
                        if t["company"] in selected_companies]
competitor_summaries = {k: v for k, v in brief.get("competitor_summaries", {}).items()
                        if k in selected_companies}
exec_summary         = brief.get("executive_summary", "")
market_trends        = brief.get("market_trends", [])

# ─────────────────────────────────────────────
# HEADER + KPIs
# ─────────────────────────────────────────────

st.markdown("# 🛡️ Competitive Intelligence Brief")
st.markdown(f"**{year_month}** — Market Monitor")

c1, c2, c3, c4 = st.columns(4)
c1.metric("Competitors", len(selected_companies))
c2.metric("Themes This Month", len(theme_summaries))
c3.metric("Avg Articles / Theme",
    round(sum(t["article_count"] for t in theme_summaries) / max(len(theme_summaries), 1), 1))
top_comp = max(
    selected_companies,
    key=lambda c: sum(t["article_count"] for t in theme_summaries if t["company"] == c),
    default="—"
)
c4.metric("Most Active Competitor", top_comp)

st.markdown("<br>", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# EXECUTIVE SUMMARY
# ─────────────────────────────────────────────

st.markdown('<div class="section-title">📋 Executive Summary</div>', unsafe_allow_html=True)
st.markdown(f'<div class="exec-box">{exec_summary}</div>', unsafe_allow_html=True)

# Market trends
if market_trends:
    st.markdown('<div class="section-title">📈 Top Market Trends</div>', unsafe_allow_html=True)
    # Row 1: first 3 trends
    row1 = market_trends[:3]
    cols1 = st.columns(len(row1))
    for i, (col, trend) in enumerate(zip(cols1, row1)):
        parts = trend.split(":", 1)
        title = parts[0].strip() if len(parts) > 1 else f"Trend {i+1}"
        body  = parts[1].strip() if len(parts) > 1 else trend
        with col:
            st.markdown(f"""
            <div class="trend-box">
                <div class="trend-title">{i+1}. {title}</div>
                <div class="trend-body">{body}</div>
            </div>
            """, unsafe_allow_html=True)
    # Row 2: remaining trends (4 and 5)
    if len(market_trends) > 3:
        st.markdown("<br>", unsafe_allow_html=True)
        row2 = market_trends[3:]
        cols2 = st.columns(len(row2))
        for i, (col, trend) in enumerate(zip(cols2, row2)):
            parts = trend.split(":", 1)
            title = parts[0].strip() if len(parts) > 1 else f"Trend {i+4}"
            body  = parts[1].strip() if len(parts) > 1 else trend
            with col:
                st.markdown(f"""
                <div class="trend-box">
                    <div class="trend-title">{i+4}. {title}</div>
                    <div class="trend-body">{body}</div>
                </div>
                """, unsafe_allow_html=True)

st.markdown("---")

# ─────────────────────────────────────────────
# ACTIVITY CHART
# ─────────────────────────────────────────────

st.markdown('<div class="section-title">📊 Theme Activity by Competitor</div>',
            unsafe_allow_html=True)

if theme_summaries:
    chart_df = pd.DataFrame([
        {"Company": t["company"], "Category": t["category"],
         "Articles": t["article_count"]}
        for t in theme_summaries
    ])
    agg = chart_df.groupby(["Company", "Category"])["Articles"].sum().reset_index()
    fig = px.bar(
        agg, x="Company", y="Articles", color="Category",
        color_discrete_sequence=["#e94560", "#0f3460", "#533483"],
    )
    fig.update_layout(
        height=260, margin=dict(l=0, r=0, t=5, b=0),
        paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
        font=dict(color="#ccc", size=11),
        legend=dict(orientation="h", y=-0.25),
        xaxis=dict(tickangle=-30),
    )
    st.plotly_chart(fig, use_container_width=True)

st.markdown("---")

# ─────────────────────────────────────────────
# COMPETITOR SECTIONS
# ─────────────────────────────────────────────

st.markdown('<div class="section-title">🏢 Competitor Analysis</div>', unsafe_allow_html=True)

for company in selected_companies:
    CATEGORY_ORDER = {"Credit Card Product": 0, "Loyalty / Program": 1, "Company News": 2}
    company_themes = sorted(
        [t for t in theme_summaries if t["company"] == company],
        key=lambda t: (CATEGORY_ORDER.get(t.get("category", ""), 3), -t.get("article_count", 0))
    )
    if not company_themes:
        continue

    total_articles = sum(t["article_count"] for t in company_themes)

    st.markdown(f'<div class="comp-header">{company}</div>', unsafe_allow_html=True)

    # Competitor summary paragraph
    comp_summary = competitor_summaries.get(company, "")
    if comp_summary:
        st.markdown(f'<div class="comp-summary">{comp_summary}</div>',
                    unsafe_allow_html=True)

    st.caption(f"📰 {len(company_themes)} themes · 📄 {total_articles} total articles")

    # Individual themes
    for t in company_themes:
        summary    = t.get("summary", "") or t.get("theme", "")
        source     = t.get("best_source", "")
        url        = t.get("best_url", "")
        count      = t.get("article_count", 0)
        date_range = t.get("date_range", "")

        source_html = (f'<a href="{url}" target="_blank" '
                       f'style="color:#6eb5ff">{source}</a>'
                       if url else source)

        st.markdown(f"""
        <div class="theme-row">
            <div class="theme-summary">{summary}</div>
            <div class="theme-meta">
                📄 {count} articles &nbsp;·&nbsp; 📅 {date_range}
                {"&nbsp;·&nbsp; 🔗 " + source_html if source else ""}
            </div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)


# ─────────────────────────────────────────────
# EXPORT TO WORD
# ─────────────────────────────────────────────

def build_docx(brief, selected_companies, theme_summaries, competitor_summaries):
    from docx import Document
    from docx.shared import Pt, RGBColor
    import io

    doc = Document()

    # Title
    h = doc.add_heading(f"Competitive Intelligence Brief — {brief.get('year_month','')}", 0)
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} · "
        f"{len(selected_companies)} competitors · {len(theme_summaries)} themes"
    )
    doc.add_paragraph("")

    # Executive Summary
    doc.add_heading("Executive Summary", 1)
    doc.add_paragraph(brief.get("executive_summary", ""))
    doc.add_paragraph("")

    # Market Trends
    trends = brief.get("market_trends", [])
    if trends:
        doc.add_heading("Top 3 Market Trends", 1)
        for i, trend in enumerate(trends, 1):
            parts = trend.split(":", 1)
            title = parts[0].strip() if len(parts) > 1 else f"Trend {i}"
            body  = parts[1].strip() if len(parts) > 1 else trend
            p = doc.add_paragraph()
            p.add_run(f"Trend {i}: {title}").bold = True
            doc.add_paragraph(body)
        doc.add_paragraph("")

    # Competitor sections
    doc.add_heading("Competitor Analysis", 1)

    for company in selected_companies:
        c_themes = [t for t in theme_summaries if t["company"] == company]
        if not c_themes:
            continue

        doc.add_heading(company, 2)

        # Competitor summary
        comp_sum = competitor_summaries.get(company, "")
        if comp_sum:
            doc.add_paragraph(comp_sum)
        doc.add_paragraph("")

        # Individual themes
        for t in sorted(c_themes, key=lambda x: x["article_count"], reverse=True):
            summary    = t.get("summary", "") or t.get("theme", "")
            source     = t.get("best_source", "")
            url        = t.get("best_url", "")
            count      = t.get("article_count", 0)
            date_range = t.get("date_range", "")

            p = doc.add_paragraph(style="List Bullet")
            p.add_run(summary)

            meta = doc.add_paragraph(style="List Bullet 2")
            meta_text = f"{count} articles · {date_range}"
            if source:
                meta_text += f" · {source}"
            if url:
                meta_text += f" · {url}"
            run = meta.add_run(meta_text)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run.font.size = Pt(9)

        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


if export_btn:
    with st.spinner("Building Word document..."):
        try:
            doc_bytes = build_docx(
                brief, selected_companies, theme_summaries, competitor_summaries
            )
            fname = f"competitive_brief_{year_month or 'latest'}.docx"
            st.sidebar.download_button(
                label="⬇️ Download .docx",
                data=doc_bytes,
                file_name=fname,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
            st.success("✅ Document ready — click Download in sidebar")
        except ImportError:
            st.error("python-docx not installed. Add it to requirements_dashboard.txt")
        except Exception as e:
            st.error(f"Export failed: {e}")


# ─────────────────────────────────────────────
# FOOTER
# ─────────────────────────────────────────────

st.markdown("---")
st.caption(
    "🛡️ Competitive Intelligence Pipeline · "
    "Google News RSS · DBSCAN Clustering · Claude AI · Streamlit"
)
